import io
import tempfile
import os
from typing import Optional
import PyPDF2
from docx import Document

class DocumentParser:
    def __init__(self):
        """Initialize the document parser."""
        pass
    
    def parse_document(self, uploaded_file) -> Optional[str]:
        """
        Parse uploaded document and extract text content.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content or None if parsing failed
        """
        if not uploaded_file:
            return None
        
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            if file_extension == 'pdf':
                return self._parse_pdf(uploaded_file)
            elif file_extension == 'docx':
                return self._parse_docx(uploaded_file)
            elif file_extension == 'txt':
                return self._parse_txt(uploaded_file)
            else:
                print(f"Unsupported file type: {file_extension}")
                return None
                
        except Exception as e:
            print(f"Error parsing document {uploaded_file.name}: {str(e)}")
            return None
    
    def _parse_pdf(self, pdf_file) -> Optional[str]:
        """
        Parse PDF file and extract text.
        
        Args:
            pdf_file: PDF file object
            
        Returns:
            Extracted text or None
        """
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
            
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text.strip())
                except Exception as e:
                    print(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
            
            if text_content:
                return "\n\n".join(text_content)
            else:
                print("No readable text found in PDF")
                return None
                
        except Exception as e:
            print(f"Error parsing PDF: {str(e)}")
            return None
    
    def _parse_docx(self, docx_file) -> Optional[str]:
        """
        Parse DOCX file and extract text.
        
        Args:
            docx_file: DOCX file object
            
        Returns:
            Extracted text or None
        """
        try:
            # Create a temporary file since python-docx needs a file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(docx_file.read())
                tmp_file_path = tmp_file.name
            
            try:
                doc = Document(tmp_file_path)
                
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                if text_content:
                    return "\n\n".join(text_content)
                else:
                    print("No readable text found in DOCX")
                    return None
                    
            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                
        except Exception as e:
            print(f"Error parsing DOCX: {str(e)}")
            return None
    
    def _parse_txt(self, txt_file) -> Optional[str]:
        """
        Parse TXT file and extract text.
        
        Args:
            txt_file: TXT file object
            
        Returns:
            Extracted text or None
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    txt_file.seek(0)
                    content = txt_file.read()
                    
                    if isinstance(content, bytes):
                        text = content.decode(encoding)
                    else:
                        text = content
                    
                    if text.strip():
                        return text.strip()
                        
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    print(f"Error with encoding {encoding}: {str(e)}")
                    continue
            
            print("Could not decode text file with any supported encoding")
            return None
            
        except Exception as e:
            print(f"Error parsing TXT: {str(e)}")
            return None
    
    def get_document_info(self, uploaded_file) -> dict:
        """
        Get basic information about the uploaded document.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary with document information
        """
        if not uploaded_file:
            return {}
        
        file_extension = uploaded_file.name.lower().split('.')[-1]
        file_size = uploaded_file.size if hasattr(uploaded_file, 'size') else len(uploaded_file.read())
        
        uploaded_file.seek(0)  # Reset after reading size
        
        info = {
            'filename': uploaded_file.name,
            'file_type': file_extension.upper(),
            'file_size': file_size,
            'file_size_mb': round(file_size / (1024 * 1024), 2)
        }
        
        # Try to get additional info based on file type
        try:
            if file_extension == 'pdf':
                uploaded_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                info['page_count'] = len(pdf_reader.pages)
                
                if pdf_reader.metadata:
                    info['title'] = pdf_reader.metadata.get('/Title', '')
                    info['author'] = pdf_reader.metadata.get('/Author', '')
                    
            elif file_extension == 'docx':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    uploaded_file.seek(0)
                    tmp_file.write(uploaded_file.read())
                    tmp_file_path = tmp_file.name
                
                try:
                    doc = Document(tmp_file_path)
                    info['paragraph_count'] = len(doc.paragraphs)
                    info['table_count'] = len(doc.tables)
                    
                    # Try to get document properties
                    if doc.core_properties:
                        info['title'] = doc.core_properties.title or ''
                        info['author'] = doc.core_properties.author or ''
                        
                finally:
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                        
        except Exception as e:
            print(f"Error getting document info: {str(e)}")
        
        return info
    
    def validate_document(self, uploaded_file) -> tuple[bool, str]:
        """
        Validate if the uploaded document can be processed.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not uploaded_file:
            return False, "No file uploaded"
        
        file_extension = uploaded_file.name.lower().split('.')[-1]
        supported_types = ['pdf', 'docx', 'txt']
        
        if file_extension not in supported_types:
            return False, f"Unsupported file type: {file_extension}. Supported types: {', '.join(supported_types)}"
        
        # Check file size (limit to 50MB)
        max_size = 50 * 1024 * 1024  # 50MB
        if hasattr(uploaded_file, 'size') and uploaded_file.size > max_size:
            return False, f"File too large. Maximum size: 50MB"
        
        # Try to parse a small portion to validate format
        try:
            uploaded_file.seek(0)
            
            if file_extension == 'pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()[:10000]))
                if len(pdf_reader.pages) == 0:
                    return False, "PDF appears to be empty or corrupted"
                    
            elif file_extension == 'docx':
                # Basic validation by trying to read file header
                header = uploaded_file.read(100)
                if not header.startswith(b'PK'):  # DOCX files are ZIP archives
                    return False, "File does not appear to be a valid DOCX document"
                    
            elif file_extension == 'txt':
                # Try to decode a portion of the file
                content = uploaded_file.read(1000)
                if isinstance(content, bytes):
                    try:
                        content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            content.decode('latin-1')
                        except UnicodeDecodeError:
                            return False, "Text file encoding not supported"
            
            uploaded_file.seek(0)  # Reset file pointer
            return True, "Document is valid"
            
        except Exception as e:
            return False, f"Error validating document: {str(e)}"
