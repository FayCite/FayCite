import streamlit as st
import os
# Streamlit Cloud automatically loads st.secrets
# No dotenv needed!

import tempfile
from io import BytesIO
from docx import Document


# Load environment variables from .env file

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from citation_processor import CitationProcessor
from pdf_library import PDFLibrary
from document_parser import DocumentParser
from citation_formatter import CitationFormatter, CitationStyle
import html
import re

def get_claim_type_color(claim_type):
    """Get color for claim type display."""
    colors = {
        'FACTUAL': '#1f77b4',  # Blue
        'STATISTICAL': '#2ca02c',  # Green  
        'THEORETICAL': '#9467bd',  # Purple
        'METHODOLOGICAL': '#ff7f0e',  # Orange
        'OPINION_INTERPRETATION': '#d62728'  # Red
    }
    return colors.get(claim_type, '#7f7f7f')

def format_claim_type_name(claim_type):
    """Format claim type for display."""
    names = {
        'FACTUAL': 'Factual',
        'STATISTICAL': 'Statistical',
        'THEORETICAL': 'Theoretical', 
        'METHODOLOGICAL': 'Methodological',
        'OPINION_INTERPRETATION': 'Opinion/Interpretation'
    }
    return names.get(claim_type, claim_type.title())

def safe_highlight_citations(text, citations):
    """Safely highlight citations in text using HTML escaping and regex word boundaries."""
    if not text or not citations:
        return html.escape(text or '')
    
    # Start with HTML-escaped text
    escaped_text = html.escape(text)
    
    # Create a set to track already processed citations to avoid duplicates
    processed_citations = set()
    
    # Sort citations by length (longest first) to avoid partial replacements
    citation_texts = []
    for citation in citations:
        citation_text = citation.get('citation', '').strip()
        if citation_text and citation_text not in processed_citations:
            citation_texts.append(citation_text)
            processed_citations.add(citation_text)
    
    # Sort by length (descending) to handle longer citations first
    citation_texts.sort(key=len, reverse=True)
    
    # Apply highlighting using regex with word boundaries for safety
    for citation_text in citation_texts:
        # Escape the citation text for use in regex
        escaped_citation = re.escape(html.escape(citation_text))
        
        # Use word boundaries to avoid partial matches, but be flexible with punctuation
        pattern = r'\b' + escaped_citation + r'\b'
        
        # Replace with highlighted version
        replacement = f'<span class="citation-highlight">{html.escape(citation_text)}</span>'
        escaped_text = re.sub(pattern, replacement, escaped_text)
    
    return escaped_text

def get_authority_badge(score):
    """Get authority score badge with appropriate color."""
    if score >= 0.8:
        return "🟢 High"
    elif score >= 0.6:
        return "🟡 Medium"
    else:
        return "🔴 Low"

# Initialize session state
if 'pdf_library' not in st.session_state:
    st.session_state.pdf_library = PDFLibrary()
if 'citation_processor' not in st.session_state:
    st.session_state.citation_processor = CitationProcessor()
if 'document_parser' not in st.session_state:
    st.session_state.document_parser = DocumentParser()
if 'citation_formatter' not in st.session_state:
    st.session_state.citation_formatter = CitationFormatter(CitationStyle.APA)
if 'citation_style' not in st.session_state:
    st.session_state.citation_style = 'APA'
if 'bibliography_parser' not in st.session_state:
    from bibliography_parser import BibliographyParser
    st.session_state.bibliography_parser = BibliographyParser()

# Check Azure OpenAI configuration and display prominent warning if not configured
if not st.session_state.citation_processor.client:
    st.error("""
    🚨 **Azure OpenAI Configuration Required**
    
    The citation processing feature requires Azure OpenAI configuration. Please add the following environment variables:
    - `AZURE_OPENAI_ENDPOINT`: Your Azure OpenAI endpoint URL
    - `AZURE_OPENAI_API_KEY`: Your Azure OpenAI API key  
    - `AZURE_OPENAI_DEPLOYMENT`: Your deployment name
    
    Without these configurations, you can still upload and manage your PDF library, but automatic citation analysis will not be available.
    """)
    st.session_state.ai_configured = False
else:
    st.session_state.ai_configured = True

st.set_page_config(
    page_title="FayCite",
    page_icon="📚",
    layout="wide"
)

st.title("📚 FayCite")
st.markdown("**Academic Citation Assistant**")
st.markdown("Automatically analyze research papers and insert citations with supporting quotes from your PDF library. Supports APA, MLA, Chicago, and IEEE citation styles.")

# Sidebar for Citation Style and Library Management
st.sidebar.header("⚙️ Citation Settings")

# Citation style selection
st.sidebar.subheader("📝 Citation Style")
citation_styles = ['APA', 'MLA', 'Chicago', 'IEEE']
selected_style = st.sidebar.selectbox(
    "Choose citation style:",
    citation_styles,
    index=citation_styles.index(st.session_state.citation_style),
    help="Select the citation style for your paper"
)

# Update citation style if changed
if selected_style != st.session_state.citation_style:
    st.session_state.citation_style = selected_style
    # Map UI names to enum values
    style_mapping = {
        'APA': CitationStyle.APA,
        'MLA': CitationStyle.MLA,
        'Chicago': CitationStyle.CHICAGO,
        'IEEE': CitationStyle.IEEE
    }
    style_enum = style_mapping.get(selected_style, CitationStyle.APA)
    st.session_state.citation_formatter.set_style(style_enum)
    st.sidebar.success(f"✅ Citation style changed to {selected_style}")

# Style information
with st.sidebar.expander("ℹ️ Citation Style Info", expanded=False):
    if selected_style == 'APA':
        st.write("**In-text:** (Author, Year, p. #)")
        st.write("**Reference list:** References")
        st.write("**Format:** Author (Year). Title.")
    elif selected_style == 'MLA':
        st.write("**In-text:** (Author Page)")
        st.write("**Reference list:** Works Cited")
        st.write("**Format:** Author. \"Title\" Year.")
    elif selected_style == 'Chicago':
        st.write("**In-text:** (Author Year, page)")
        st.write("**Reference list:** Bibliography")
        st.write("**Format:** Author. Year. \"Title\".")
    elif selected_style == 'IEEE':
        st.write("**In-text:** [Number]")
        st.write("**Reference list:** References")
        st.write("**Format:** [#] Author, \"Title\" Year.")

st.sidebar.header("📚 Reference Library Management")

# Bibliography upload section
st.sidebar.subheader("📋 Zotero Bibliography")
uploaded_bib = st.sidebar.file_uploader(
    "Upload Zotero bibliography (.txt)",
    type=['txt'],
    help="Export your Zotero library as a bibliography in .txt format"
)

if uploaded_bib:
    content = uploaded_bib.read().decode('utf-8')
    if st.session_state.bibliography_parser.parse_zotero_txt(content):
        stats = st.session_state.bibliography_parser.get_statistics()
        st.sidebar.success(f"✅ Bibliography loaded: {stats['total_entries']} entries")
        
        with st.sidebar.expander("Bibliography Stats", expanded=False):
            st.write(f"**Total entries:** {stats['total_entries']}")
            st.write("**Entry types:**")
            for entry_type, count in stats['types'].items():
                st.write(f"  • {entry_type}: {count}")
            st.write(f"**With DOI:** {stats['entries_with_doi']}")
            st.write(f"**With year:** {stats['entries_with_year']}")
    else:
        st.sidebar.error("Failed to parse bibliography file")

st.sidebar.subheader("📄 PDF Sources")

uploaded_pdfs = st.sidebar.file_uploader(
    "Upload PDF sources",
    type=['pdf'],
    accept_multiple_files=True,
    help="Upload academic papers, books, and other sources for citation"
)

if uploaded_pdfs:
    for pdf_file in uploaded_pdfs:
        if pdf_file.name not in st.session_state.pdf_library.get_library_files():
            with st.spinner(f"Processing {pdf_file.name}..."):
                success = st.session_state.pdf_library.add_pdf(pdf_file)
                if success:
                    # Check if PDF matches bibliography entry
                    pdf_content = st.session_state.pdf_library.get_pdf_content(pdf_file.name)
                    if pdf_content and hasattr(st.session_state, 'bibliography_parser') and len(st.session_state.bibliography_parser.get_all_entries()) > 0:
                        matching_entry = st.session_state.bibliography_parser.find_matching_entry(
                            pdf_file.name, 
                            pdf_content.get('metadata', {})
                        )
                        if matching_entry:
                            st.sidebar.success(f"✅ {pdf_file.name} (Validated against bibliography)")
                        else:
                            st.sidebar.warning(f"⚠️ {pdf_file.name} (Not found in bibliography)")
                    else:
                        st.sidebar.success(f"Added: {pdf_file.name}")
                else:
                    st.sidebar.error(f"Failed to process: {pdf_file.name}")

# Display current library with validation status
library_files = st.session_state.pdf_library.get_library_files()
if library_files:
    st.sidebar.subheader("Current Library")
    for filename in library_files:
        col1, col2 = st.sidebar.columns([3, 1])
        
        # Check validation status
        validation_icon = "📄"
        if hasattr(st.session_state, 'bibliography_parser') and len(st.session_state.bibliography_parser.get_all_entries()) > 0:
            pdf_content = st.session_state.pdf_library.get_pdf_content(filename)
            if pdf_content:
                matching_entry = st.session_state.bibliography_parser.find_matching_entry(
                    filename, 
                    pdf_content.get('metadata', {})
                )
                validation_icon = "✅" if matching_entry else "⚠️"
        
        col1.write(f"{validation_icon} {filename}")
        if col2.button("🗑️", key=f"delete_{filename}", help="Remove from library"):
            st.session_state.pdf_library.remove_pdf(filename)
            st.rerun()

# Main interface tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs(["📝 Process Paper", "🔍 Results", "📋 References", "✅ Validation", "📄 Document Viewer"])

with tab1:
    st.header("Upload Research Paper")
    
    uploaded_paper = st.file_uploader(
        "Choose your research paper",
        type=['pdf', 'docx', 'txt'],
        help="Upload your draft research paper for citation analysis"
    )
    
    if uploaded_paper:
        st.success(f"Paper uploaded: {uploaded_paper.name}")
        
        # Parse document
        with st.spinner("Parsing document..."):
            paper_content = st.session_state.document_parser.parse_document(uploaded_paper)
        
        if paper_content:
            st.subheader("Document Preview")
            with st.expander("View original text", expanded=False):
                st.text_area("Original Content", paper_content, height=200, disabled=True)
            
            # Process button with configuration check
            button_disabled = not library_files or not st.session_state.get('ai_configured', False)
            if st.button("🔄 Analyze and Insert Citations", type="primary", disabled=button_disabled):
                if not library_files:
                    st.error("Please upload PDF sources to your library first.")
                elif not st.session_state.get('ai_configured', False):
                    st.error("Azure OpenAI configuration required. Please check the configuration warning above.")
                else:
                    with st.spinner("Analyzing paper and finding citations..."):
                        # Process citations with bibliography validation and selected citation style
                        result = st.session_state.citation_processor.process_paper(
                            paper_content, 
                            st.session_state.pdf_library,
                            st.session_state.bibliography_parser if hasattr(st.session_state, 'bibliography_parser') else None,
                            st.session_state.citation_formatter
                        )
                        
                        if result and 'error' not in result:
                            st.session_state.processed_result = result
                            st.success("✅ Citation analysis complete! Check the Results tab.")
                            st.rerun()
                        elif result and 'error' in result:
                            st.error(f"Configuration error: {result['error']}")
                        else:
                            st.error("Failed to process the paper. Please check your Azure OpenAI configuration and try again.")
            
            if not library_files:
                st.warning("⚠️ No PDF sources in library. Please upload sources using the sidebar.")

with tab2:
    st.header("Citation Results")
    
    if 'processed_result' in st.session_state:
        result = st.session_state.processed_result
        
        # Before/After comparison
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Original Text")
            st.text_area("Before", result['original_text'], height=400, disabled=True)
        
        with col2:
            st.subheader("📚 With Citations")
            st.text_area("After", result['cited_text'], height=400, disabled=True)
        
        # Enhanced Statistics with Claim Type Analysis
        st.subheader("📊 Citation Statistics")
        stats_col1, stats_col2, stats_col3, stats_col4 = st.columns(4)
        
        with stats_col1:
            st.metric("Claims Identified", result['stats']['claims_identified'])
        
        with stats_col2:
            st.metric("Citations Added", result['stats']['citations_added'])
        
        with stats_col3:
            st.metric("Sources Used", result['stats']['sources_used'])
        
        with stats_col4:
            avg_authority = sum(citation.get('source_authority_score', 0.5) for citation in result['citations']) / len(result['citations']) if result['citations'] else 0
            st.metric("Avg Source Authority", f"{avg_authority:.2f}")
        
        # Claim Type Distribution
        if result['citations']:
            st.subheader("🎯 Claim Type Analysis")
            
            # Count claim types
            claim_type_counts = {}
            for citation in result['citations']:
                claim_type = citation.get('claim_type', 'FACTUAL')
                claim_type_counts[claim_type] = claim_type_counts.get(claim_type, 0) + 1
            
            # Display claim type distribution
            type_col1, type_col2 = st.columns([1, 2])
            
            with type_col1:
                st.write("**Distribution:**")
                for claim_type, count in claim_type_counts.items():
                    percentage = (count / len(result['citations'])) * 100
                    color = get_claim_type_color(claim_type)
                    st.markdown(f"<span style='color: {color}'>●</span> **{format_claim_type_name(claim_type)}**: {count} ({percentage:.1f}%)", unsafe_allow_html=True)
            
            with type_col2:
                st.write("**Claim Type Strategies:**")
                with st.expander("📘 Claim Type Explanations", expanded=False):
                    st.markdown("""
                    **🔵 Factual Claims**: Require authoritative, primary sources (gov data, established references)
                    
                    **🟢 Statistical Claims**: Require data sources with clear methodology (recent studies, sample sizes)
                    
                    **🟣 Theoretical Claims**: Require foundational academic sources (seminal works, established theories)
                    
                    **🟠 Methodological Claims**: Require peer-reviewed research methods (established protocols)
                    
                    **🔴 Opinion/Interpretation Claims**: Require balanced perspectives (multiple viewpoints, recent analysis)
                    """)
        
        # Enhanced Detailed citations with categorization
        if result['citations']:
            st.subheader("🔗 Enhanced Citation Details")
            for i, citation in enumerate(result['citations'], 1):
                claim_type = citation.get('claim_type', 'FACTUAL')
                authority_score = citation.get('source_authority_score', 0.5)
                color = get_claim_type_color(claim_type)
                authority_badge = get_authority_badge(authority_score)
                
                # Create colored title with claim type and authority
                title_html = f"<span style='color: {color}'>●</span> Citation {i}: {citation['source']} | {format_claim_type_name(claim_type)} | {authority_badge}"
                
                with st.expander(f"Citation {i}: {citation['source']}", expanded=False):
                    st.markdown(title_html, unsafe_allow_html=True)
                    
                    # Claim with type information
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.write(f"**Claim:** {citation['claim']}")
                    with col2:
                        st.markdown(f"<span style='color: {color}'>**{format_claim_type_name(claim_type)}**</span>", unsafe_allow_html=True)
                    
                    # Claim reasoning if available
                    if citation.get('claim_reasoning'):
                        st.write(f"**Categorization Reason:** {citation['claim_reasoning']}")
                    
                    st.write(f"**Supporting Quote:** \"{citation['quote']}\"")
                    st.write(f"**Page:** {citation['page']}")
                    st.write(f"**Citation:** {citation['citation']}")
                    
                    # Enhanced metadata display
                    metrics_col1, metrics_col2, metrics_col3 = st.columns(3)
                    with metrics_col1:
                        st.metric("Authority Score", f"{authority_score:.2f}")
                    with metrics_col2:
                        claim_match = citation.get('claim_type_match', 'Not evaluated')
                        match_score = "Strong" if "Strong" in claim_match else "Good" if "Good" in claim_match else "Moderate" if "Moderate" in claim_match else "Weak"
                        st.metric("Type Match", match_score)
                    with metrics_col3:
                        # Calculate relevance indicator
                        relevance = "High" if authority_score > 0.7 and "Strong" in claim_match else "Medium" if authority_score > 0.5 else "Low"
                        st.metric("Overall Quality", relevance)
                    
                    # Source appropriateness explanation
                    if citation.get('claim_type_match'):
                        st.info(f"💡 **Source Assessment:** {citation['claim_type_match']}")
                    
                    # Citation strategy explanation
                    strategy_explanations = {
                        'FACTUAL': "This factual claim requires authoritative, primary sources for credibility.",
                        'STATISTICAL': "This statistical claim needs data sources with clear methodology and sample information.",
                        'THEORETICAL': "This theoretical claim benefits from foundational academic sources and established frameworks.",
                        'METHODOLOGICAL': "This methodological claim requires peer-reviewed research methods and established protocols.",
                        'OPINION_INTERPRETATION': "This interpretive claim benefits from balanced perspectives and recent analytical sources."
                    }
                    
                    if claim_type in strategy_explanations:
                        st.caption(f"📘 {strategy_explanations[claim_type]}")
        
        # Download options
        st.subheader("💾 Download Results")
        col1, col2 = st.columns(2)
        
        with col1:
            st.download_button(
                "📄 Download Cited Paper",
                result['cited_text'],
                file_name="paper_with_citations.txt",
                mime="text/plain"
            )
        
        with col2:
            if result['references']:
                st.download_button(
                    "📋 Download References",
                    result['references'],
                    file_name="references.txt",
                    mime="text/plain"
                )
    
    else:
        st.info("No results yet. Please process a paper in the 'Process Paper' tab first.")

with tab3:
    st.header("References List")
    
    if 'processed_result' in st.session_state and st.session_state.processed_result.get('references'):
        current_style = st.session_state.citation_style
        style_headers = {
            'APA': 'References',
            'MLA': 'Works Cited', 
            'Chicago': 'Bibliography',
            'IEEE': 'References'
        }
        header_text = style_headers.get(current_style, 'References')
        st.subheader(f"📋 {header_text} ({current_style} Style)")
        st.text_area(
            "References",
            st.session_state.processed_result['references'],
            height=300,
            disabled=True
        )
        
        # Copy to clipboard helper
        st.code(st.session_state.processed_result['references'], language=None)
        st.caption("💡 Tip: You can copy the references from the code block above")
    
    else:
        st.info("No references generated yet. Process a paper first to see the references list.")

with tab4:
    st.header("PDF Validation Status")
    
    if hasattr(st.session_state, 'bibliography_parser') and len(st.session_state.bibliography_parser.get_all_entries()) > 0:
        if library_files:
            st.subheader("📊 Validation Summary")
            
            validated_count = 0
            unvalidated_files = []
            validated_files = []
            
            for filename in library_files:
                pdf_content = st.session_state.pdf_library.get_pdf_content(filename)
                if pdf_content:
                    matching_entry = st.session_state.bibliography_parser.find_matching_entry(
                        filename, 
                        pdf_content.get('metadata', {})
                    )
                    if matching_entry:
                        validated_count += 1
                        validated_files.append((filename, matching_entry))
                    else:
                        unvalidated_files.append(filename)
            
            # Display metrics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total PDFs", len(library_files))
            with col2:
                st.metric("Validated", validated_count)
            with col3:
                st.metric("Unvalidated", len(unvalidated_files))
            
            # Show validated files
            if validated_files:
                st.subheader("✅ Validated PDFs")
                for filename, bib_entry in validated_files:
                    with st.expander(f"✅ {filename}", expanded=False):
                        st.write(f"**Title:** {bib_entry.get('title', 'N/A')}")
                        authors = bib_entry.get('authors', [])
                        if authors:
                            st.write(f"**Authors:** {', '.join(authors)}")
                        if bib_entry.get('year'):
                            st.write(f"**Year:** {bib_entry['year']}")
                        if bib_entry.get('journal'):
                            st.write(f"**Journal:** {bib_entry['journal']}")
                        if bib_entry.get('doi'):
                            st.write(f"**DOI:** {bib_entry['doi']}")
            
            # Show unvalidated files
            if unvalidated_files:
                st.subheader("⚠️ Unvalidated PDFs")
                st.warning("These PDFs could not be matched with entries in your Zotero bibliography:")
                for filename in unvalidated_files:
                    st.write(f"⚠️ {filename}")
                
                st.info("💡 **Tip:** Make sure these PDFs are properly referenced in your Zotero library and exported in the bibliography file.")
        else:
            st.info("No PDFs uploaded yet. Upload some PDFs to see validation status.")
    else:
        st.info("Upload a Zotero bibliography file to see PDF validation status.")

def create_word_document(content, citations, references, style_name):
    """Create a Word document with the cited content."""
    doc = Document()
    
    # Set document style (skip if not available)
    try:
        for paragraph in doc.paragraphs:
            if paragraph.runs:
                paragraph.runs[0].font.name = 'Times New Roman'
                paragraph.runs[0].font.size = Pt(12)
    except (AttributeError, KeyError):
        pass
    
    # Add title
    title = doc.add_heading('Research Paper with Citations', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add citation style info
    style_para = doc.add_paragraph(f'Citation Style: {style_name}')
    style_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line
    
    # Add main content
    content_para = doc.add_paragraph(content)
    
    # Add page break
    doc.add_page_break()
    
    # Add references section
    if references:
        ref_heading_map = {
            'APA': 'References',
            'MLA': 'Works Cited', 
            'Chicago': 'Bibliography',
            'IEEE': 'References'
        }
        ref_title = ref_heading_map.get(style_name, 'References')
        
        doc.add_heading(ref_title, 1)
        
        # Add each reference
        for ref_line in references.split('\n'):
            if ref_line.strip():
                doc.add_paragraph(ref_line.strip())
    
    return doc

with tab5:
    st.header("📄 Document Viewer")
    
    if 'processed_result' in st.session_state and st.session_state.processed_result:
        result = st.session_state.processed_result
        
        # Document viewer header with download button
        col1, col2 = st.columns([3, 1])
        with col1:
            st.subheader("Your Paper with Citations")
        with col2:
            # Create Word document
            word_doc = create_word_document(
                result['cited_text'],
                result['citations'],
                result.get('references', ''),
                st.session_state.citation_style
            )
            
            # Save to bytes
            doc_buffer = BytesIO()
            word_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.download_button(
                label="📥 Download as Word",
                data=doc_buffer.getvalue(),
                file_name=f"cited_paper_{st.session_state.citation_style.lower()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                help="Download your paper with citations as a Word document"
            )
        
        st.markdown("---")
        
        # Grammarly-style document display
        st.markdown("""
        <style>
        .document-viewer {
            background-color: white;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 2rem;
            margin: 1rem 0;
            font-family: 'Times New Roman', serif;
            font-size: 16px;
            line-height: 1.6;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .citation-highlight {
            background-color: #fff3cd;
            border-left: 3px solid #ffc107;
            padding: 2px 4px;
            margin: 0 2px;
            border-radius: 3px;
        }
        .document-title {
            text-align: center;
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 1rem;
            color: #333;
        }
        .citation-style-info {
            text-align: center;
            color: #666;
            font-style: italic;
            margin-bottom: 2rem;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Document content with safely highlighted citations
        cited_text = result['cited_text']
        
        # Create safely highlighted version using HTML escaping and regex
        display_text = safe_highlight_citations(cited_text, result.get('citations', []))
        
        # Display the document in a clean format with sanitized content
        # Note: Only CSS styles and our controlled highlight spans are allowed
        safe_style_name = html.escape(st.session_state.citation_style)
        st.markdown(f"""
        <div class="document-viewer">
            <div class="document-title">Research Paper with Citations</div>
            <div class="citation-style-info">Citation Style: {safe_style_name}</div>
            <div>{display_text}</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Citation summary
        if 'citations' in result and result['citations']:
            st.subheader("📊 Citation Summary")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Total Citations", len(result['citations']))
            
            with col2:
                if 'stats' in result and 'avg_source_authority' in result['stats']:
                    authority_score = result['stats']['avg_source_authority']
                    st.metric("Avg Source Authority", f"{authority_score:.1f}/1.0")
                else:
                    st.metric("Avg Source Authority", "N/A")
            
            with col3:
                if 'stats' in result and 'claim_type_distribution' in result['stats']:
                    most_common = max(result['stats']['claim_type_distribution'].items(), key=lambda x: x[1])
                    st.metric("Most Common Claim", format_claim_type_name(most_common[0]))
        
        # Quick actions
        st.subheader("🔧 Quick Actions")
        action_col1, action_col2, action_col3 = st.columns(3)
        
        with action_col1:
            if st.button("📋 Copy Text", help="Copy the cited text to clipboard"):
                st.code(result['cited_text'], language=None)
                st.success("Text ready to copy from code block above!")
        
        with action_col2:
            if st.button("🔍 View Detailed Results", help="Go to detailed results tab"):
                st.info("Switch to the 'Results' tab for detailed citation analysis")
                
        with action_col3:
            if st.button("📚 View References", help="Go to references tab"):
                st.info("Switch to the 'References' tab for the complete reference list")
        
        # Document statistics
        st.subheader("📈 Document Statistics")
        stats_col1, stats_col2, stats_col3, stats_col4 = st.columns(4)
        
        with stats_col1:
            word_count = len(result['cited_text'].split())
            st.metric("Word Count", f"{word_count:,}")
        
        with stats_col2:
            char_count = len(result['cited_text'])
            st.metric("Character Count", f"{char_count:,}")
        
        with stats_col3:
            paragraph_count = len([p for p in result['cited_text'].split('\n\n') if p.strip()])
            st.metric("Paragraphs", paragraph_count)
        
        with stats_col4:
            if 'citations' in result:
                citation_density = len(result['citations']) / max(word_count / 100, 1)  # Citations per 100 words
                st.metric("Citation Density", f"{citation_density:.1f}/100 words")
            else:
                st.metric("Citation Density", "0/100 words")
                
    else:
        st.info("No processed document available. Process a paper first in the 'Process Paper' tab to view it here.")
        
        # Show example of what the viewer will look like
        st.subheader("📋 Preview")
        st.markdown("This is what your document viewer will look like:")
        
        st.markdown("""
        <div class="document-viewer">
            <div class="document-title">Sample Research Paper</div>
            <div class="citation-style-info">Citation Style: APA</div>
            <div>
                This is a sample paragraph of your research paper. When you process a document, 
                citations will be automatically inserted and highlighted like this: 
                <span class="citation-highlight">(Smith, 2023, p. 42)</span>. 
                The document will be displayed in a clean, professional format similar to 
                Grammarly's document editor.
            </div>
        </div>
        """, unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown(
    "💡 **Tips:** "
    "• Upload multiple PDF sources for better citation coverage "
    "• Ensure your research paper has clear claims and statements "
    "• Review generated citations for accuracy before final submission"
)
